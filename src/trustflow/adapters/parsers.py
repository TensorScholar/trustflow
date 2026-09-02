"""Questionnaire parsers."""

from __future__ import annotations

import csv
import json
from collections.abc import Iterable
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.table import Table
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from pypdf import PdfReader

from trustflow.adapters.safety import file_sha256, inspect_document
from trustflow.domain.classification import classify_sensitivity
from trustflow.domain.errors import InvalidQuestionnaireError, UnsupportedFormatError
from trustflow.domain.models import (
    DocumentFormat,
    PolicySettings,
    Question,
    QuestionLocation,
    Questionnaire,
)


def _question(identifier: str, text: str, location: QuestionLocation) -> Question:
    return Question(
        id=identifier,
        text=text.strip(),
        location=location,
        sensitivity=classify_sensitivity(text),
    )


def _iter_docx_table_paragraphs(
    table: Table,
    path: str,
    seen_cells: set[int],
) -> Iterable[tuple[str, int, str]]:
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            marker = id(cell._tc)
            if marker in seen_cells:
                continue
            seen_cells.add(marker)
            cell_path = f"{path}/row:{row_index}/cell:{cell_index}"
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield paragraph.text, paragraph_index, cell_path
            for nested_index, nested in enumerate(cell.tables):
                yield from _iter_docx_table_paragraphs(
                    nested,
                    f"{cell_path}/table:{nested_index}",
                    seen_cells,
                )


def _xlsx_column_hidden(sheet: Worksheet, column: int) -> bool:
    for dimension in sheet.column_dimensions.values():
        if not dimension.hidden:
            continue
        minimum = dimension.min or 0
        maximum = dimension.max or minimum
        if minimum <= column <= maximum:
            return True
    return False


class ParserRegistry:
    def __init__(self, policy: PolicySettings | None = None) -> None:
        self.policy = policy or PolicySettings()

    def parse(self, path: Path) -> Questionnaire:
        fmt = inspect_document(path, self.policy)
        source_digest = file_sha256(path)
        if fmt is DocumentFormat.XLSX:
            questions = self._xlsx(path)
        elif fmt is DocumentFormat.DOCX:
            questions = self._docx(path)
        elif fmt is DocumentFormat.CSV:
            questions = self._csv(path)
        elif fmt is DocumentFormat.JSON:
            questions = self._json(path)
        elif fmt is DocumentFormat.MARKDOWN:
            questions = self._markdown(path)
        elif fmt is DocumentFormat.PDF:
            questions = self._pdf(path)
        else:
            raise UnsupportedFormatError(fmt.value)
        if file_sha256(path) != source_digest:
            raise InvalidQuestionnaireError("questionnaire changed during import")
        if not questions:
            raise InvalidQuestionnaireError("questionnaire contains no detectable questions")
        if len(questions) > self.policy.maximum_questions:
            raise InvalidQuestionnaireError("questionnaire exceeds configured question limit")
        return Questionnaire(
            id=f"qnr_{uuid4().hex}",
            title=path.stem,
            source_path=str(path.resolve()),
            source_digest=source_digest,
            format=fmt,
            questions=tuple(questions),
        )

    def _xlsx(self, path: Path) -> list[Question]:
        workbook = load_workbook(path, read_only=False, data_only=False, keep_links=False)
        try:
            questions: list[Question] = []
            index = 1
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        value = cell.value
                        if not (isinstance(value, str) and value.strip().endswith("?")):
                            continue
                        if sheet.sheet_state != "visible":
                            raise InvalidQuestionnaireError(
                                f"question detected on hidden XLSX sheet: {sheet.title}"
                            )
                        row_dimension = sheet.row_dimensions.get(cell.row)
                        if row_dimension is not None and row_dimension.hidden:
                            raise InvalidQuestionnaireError(
                                "question detected on hidden XLSX row: "
                                f"{sheet.title}!{cell.coordinate}"
                            )
                        if _xlsx_column_hidden(sheet, cell.column):
                            raise InvalidQuestionnaireError(
                                f"question detected in hidden XLSX column: "
                                f"{sheet.title}!{cell.coordinate}"
                            )
                        questions.append(
                            _question(
                                f"q{index}",
                                value,
                                QuestionLocation(
                                    format=DocumentFormat.XLSX,
                                    sheet=sheet.title,
                                    cell=cell.coordinate,
                                ),
                            )
                        )
                        index += 1
            return questions
        finally:
            workbook.close()

    def _docx(self, path: Path) -> list[Question]:
        document = Document(str(path))
        questions: list[Question] = []
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if text.endswith("?"):
                questions.append(
                    _question(
                        f"q{len(questions) + 1}",
                        text,
                        QuestionLocation(format=DocumentFormat.DOCX, paragraph=index),
                    )
                )
        seen_cells: set[int] = set()
        for table_index, table in enumerate(document.tables):
            for text, paragraph_index, cell_path in _iter_docx_table_paragraphs(
                table,
                f"table:{table_index}",
                seen_cells,
            ):
                cleaned = text.strip()
                if cleaned.endswith("?"):
                    questions.append(
                        _question(
                            f"q{len(questions) + 1}",
                            cleaned,
                            QuestionLocation(
                                format=DocumentFormat.DOCX,
                                key=f"{cell_path}/paragraph:{paragraph_index}",
                            ),
                        )
                    )
        return questions

    def _csv(self, path: Path) -> list[Question]:
        questions: list[Question] = []
        with path.open(newline="", encoding="utf-8-sig") as handle:
            for row_index, row in enumerate(csv.reader(handle), start=1):
                for cell_index, value in enumerate(row):
                    if value.strip().endswith("?"):
                        questions.append(
                            _question(
                                f"q{len(questions) + 1}",
                                value,
                                QuestionLocation(
                                    format=DocumentFormat.CSV,
                                    row=row_index,
                                    key=str(cell_index),
                                ),
                            )
                        )
        return questions

    def _json(self, path: Path) -> list[Question]:
        payload = json.loads(path.read_text(encoding="utf-8"))
        rows = (
            payload["questions"]
            if isinstance(payload, dict) and "questions" in payload
            else payload
        )
        if not isinstance(rows, list):
            raise InvalidQuestionnaireError("JSON questionnaire must contain a list of questions")
        questions = []
        for index, row in enumerate(rows):
            text = row.get("question") if isinstance(row, dict) else row
            if not isinstance(text, str):
                raise InvalidQuestionnaireError("question must be a string")
            questions.append(
                _question(
                    f"q{index + 1}",
                    text,
                    QuestionLocation(format=DocumentFormat.JSON, key=str(index)),
                )
            )
        return questions

    def _markdown(self, path: Path) -> list[Question]:
        questions: list[Question] = []
        for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
            text = line.lstrip("#-* 0123456789.").strip()
            if text.endswith("?"):
                questions.append(
                    _question(
                        f"q{len(questions) + 1}",
                        text,
                        QuestionLocation(format=DocumentFormat.MARKDOWN, row=line_number),
                    )
                )
        return questions

    def _pdf(self, path: Path) -> list[Question]:
        reader = PdfReader(path)
        if len(reader.pages) > self.policy.maximum_pdf_pages:
            raise InvalidQuestionnaireError("PDF exceeds configured page limit")
        questions: list[Question] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for line in text.splitlines():
                cleaned = line.strip()
                if cleaned.endswith("?"):
                    questions.append(
                        _question(
                            f"q{len(questions) + 1}",
                            cleaned,
                            QuestionLocation(format=DocumentFormat.PDF, row=page_number),
                        )
                    )
        return questions
