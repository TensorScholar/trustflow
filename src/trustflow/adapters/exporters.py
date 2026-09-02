"""Safe questionnaire exporters."""

from __future__ import annotations

import csv
import json
import os
import shutil
import tempfile
from collections.abc import Iterator
from contextlib import contextmanager
from pathlib import Path
from typing import cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.worksheet.worksheet import Worksheet

from trustflow.adapters.safety import file_sha256, neutralize_spreadsheet_formula
from trustflow.domain.errors import (
    InvalidTransitionError,
    UnsafeExportError,
    UnsupportedFormatError,
)
from trustflow.domain.models import (
    AnswerStatus,
    DocumentFormat,
    DraftAnswer,
    ExportResult,
    Questionnaire,
    ReviewDecision,
    ReviewState,
)
from trustflow.domain.review import review_binding_error

_MISSING_DIGEST = "0" * 64
_MAX_XLSX_COLUMN = 16_384


def _final_text(answer: DraftAnswer, review: ReviewDecision | None) -> str:
    if not answer.evidence:
        raise InvalidTransitionError("claim without evidence cannot be exported")
    if any(item.source_digest == _MISSING_DIGEST for item in answer.evidence):
        raise InvalidTransitionError("claim evidence snapshot is missing a source fingerprint")
    if review is None:
        if answer.status is not AnswerStatus.ANSWERED:
            raise InvalidTransitionError("unresolved answer cannot be exported")
        return answer.text
    binding_error = review_binding_error(answer, review)
    if binding_error is not None:
        raise InvalidTransitionError(f"review does not bind to current answer: {binding_error}")
    if review.state not in {ReviewState.APPROVED, ReviewState.EDITED}:
        raise InvalidTransitionError("rejected review cannot be exported")
    if answer.status is AnswerStatus.UNANSWERABLE:
        raise InvalidTransitionError("unanswerable answer cannot become an external claim")
    return review.final_text


def _answer_mapping(
    questionnaire: Questionnaire,
    answers: list[DraftAnswer],
) -> dict[str, DraftAnswer]:
    mapping: dict[str, DraftAnswer] = {}
    for answer in answers:
        if answer.questionnaire_id != questionnaire.id:
            raise InvalidTransitionError("answer belongs to a different questionnaire")
        if answer.question_id in mapping:
            raise InvalidTransitionError(f"duplicate answer for question: {answer.question_id}")
        mapping[answer.question_id] = answer
    expected = {question.id for question in questionnaire.questions}
    actual = set(mapping)
    if actual != expected:
        missing = sorted(expected - actual)
        extra = sorted(actual - expected)
        raise InvalidTransitionError(f"answer set mismatch; missing={missing}, extra={extra}")
    return mapping


def _verify_source_identity(questionnaire: Questionnaire, source: Path) -> None:
    if questionnaire.source_digest == _MISSING_DIGEST:
        raise UnsafeExportError("questionnaire source fingerprint is missing")
    if file_sha256(source) != questionnaire.source_digest:
        raise UnsafeExportError("questionnaire source changed since import")


def _source_and_destination(questionnaire: Questionnaire, output: Path) -> tuple[Path, Path]:
    try:
        source = Path(questionnaire.source_path).resolve(strict=True)
    except FileNotFoundError as exc:
        raise UnsafeExportError("questionnaire source file no longer exists") from exc
    destination = output.expanduser().resolve(strict=False)
    if source == destination:
        raise UnsafeExportError("export destination must differ from the source document")
    if destination.exists():
        if destination.is_dir():
            raise UnsafeExportError("export destination is a directory")
        raise UnsafeExportError("export destination already exists")
    _verify_source_identity(questionnaire, source)
    destination.parent.mkdir(parents=True, exist_ok=True)
    return source, destination


@contextmanager
def _atomic_destination(destination: Path) -> Iterator[Path]:
    descriptor, raw_path = tempfile.mkstemp(
        dir=destination.parent,
        prefix=f".{destination.name}.",
        suffix=f".tmp{destination.suffix}",
    )
    os.close(descriptor)
    temporary = Path(raw_path)
    try:
        yield temporary
        try:
            # A same-directory hard link is an atomic create-if-absent commit. Unlike
            # os.replace(), it cannot overwrite a destination created concurrently.
            os.link(temporary, destination)
        except FileExistsError as exc:
            raise UnsafeExportError("export destination was created concurrently") from exc
        except OSError as exc:
            raise UnsafeExportError("atomic no-overwrite export commit failed") from exc
        temporary.unlink()
    finally:
        temporary.unlink(missing_ok=True)


def _docx_paragraph(document: DocumentObject, key: str) -> Paragraph:
    table = None
    cell = None
    row_index: int | None = None
    for segment in key.split("/"):
        kind, separator, raw_index = segment.partition(":")
        if not separator or not raw_index.isdigit():
            raise UnsafeExportError(f"invalid DOCX location: {key}")
        index = int(raw_index)
        if kind == "table":
            tables = document.tables if cell is None else cell.tables
            try:
                table = tables[index]
            except IndexError as exc:
                raise UnsafeExportError(f"DOCX table location no longer exists: {key}") from exc
            cell = None
            row_index = None
        elif kind == "row":
            if table is None:
                raise UnsafeExportError(f"invalid DOCX row location: {key}")
            row_index = index
        elif kind == "cell":
            if table is None or row_index is None:
                raise UnsafeExportError(f"invalid DOCX cell location: {key}")
            try:
                cell = table.rows[row_index].cells[index]
            except IndexError as exc:
                raise UnsafeExportError(f"DOCX cell location no longer exists: {key}") from exc
        elif kind == "paragraph":
            if cell is None:
                raise UnsafeExportError(f"invalid DOCX paragraph location: {key}")
            try:
                return cast(Paragraph, cell.paragraphs[index])
            except IndexError as exc:
                raise UnsafeExportError(f"DOCX paragraph location no longer exists: {key}") from exc
        else:
            raise UnsafeExportError(f"unknown DOCX location segment: {kind}")
    raise UnsafeExportError(f"DOCX location does not identify a paragraph: {key}")


def _xlsx_target(sheet: Worksheet, question_cell: str) -> Cell:
    cell = sheet[question_cell]
    if not isinstance(cell, Cell):
        raise UnsafeExportError("XLSX question location is not a writable cell")
    merged_question = next(
        (merged for merged in sheet.merged_cells.ranges if cell.coordinate in merged),
        None,
    )
    target_column = merged_question.max_col + 1 if merged_question is not None else cell.column + 1
    if target_column > _MAX_XLSX_COLUMN:
        raise UnsafeExportError("XLSX question has no writable adjacent answer column")
    target = sheet.cell(row=cell.row, column=target_column)
    if any(target.coordinate in merged for merged in sheet.merged_cells.ranges):
        raise UnsafeExportError(
            f"XLSX answer target is part of a merged range: {sheet.title}!{target.coordinate}"
        )
    if not isinstance(target, Cell):
        raise UnsafeExportError("XLSX answer target is not a writable cell")
    if target.value not in (None, ""):
        raise UnsafeExportError(
            f"XLSX answer target is occupied: {sheet.title}!{target.coordinate}"
        )
    return target


class ExporterRegistry:
    def export(
        self,
        questionnaire: Questionnaire,
        answers: list[DraftAnswer],
        reviews: dict[str, ReviewDecision],
        output: Path,
    ) -> ExportResult:
        mapping = _answer_mapping(questionnaire, answers)
        source, destination = _source_and_destination(questionnaire, output)
        if questionnaire.format is DocumentFormat.XLSX:
            self._xlsx(questionnaire, mapping, reviews, source, destination)
        elif questionnaire.format is DocumentFormat.DOCX:
            self._docx(questionnaire, mapping, reviews, source, destination)
        elif questionnaire.format is DocumentFormat.CSV:
            self._csv(questionnaire, mapping, reviews, source, destination)
        elif questionnaire.format in {
            DocumentFormat.JSON,
            DocumentFormat.MARKDOWN,
            DocumentFormat.PDF,
        }:
            self._json(questionnaire, mapping, reviews, source, destination)
        else:
            raise UnsupportedFormatError(questionnaire.format.value)
        return ExportResult(
            questionnaire_id=questionnaire.id,
            output_path=str(destination),
            format=questionnaire.format,
            answered=sum(item.status is AnswerStatus.ANSWERED for item in answers),
            review_required=sum(
                item.status in {AnswerStatus.REVIEW_REQUIRED, AnswerStatus.CONFLICT}
                for item in answers
            ),
            unanswerable=sum(
                item.status in {AnswerStatus.UNANSWERABLE, AnswerStatus.STALE} for item in answers
            ),
        )

    def _xlsx(
        self,
        questionnaire: Questionnaire,
        answers: dict[str, DraftAnswer],
        reviews: dict[str, ReviewDecision],
        source: Path,
        destination: Path,
    ) -> None:
        with _atomic_destination(destination) as temporary:
            shutil.copyfile(source, temporary)
            _verify_source_identity(questionnaire, temporary)
            workbook = load_workbook(temporary, keep_links=False)
            try:
                for question in questionnaire.questions:
                    answer = answers[question.id]
                    sheet_name = question.location.sheet
                    cell_name = question.location.cell
                    if not sheet_name or not cell_name:
                        raise UnsafeExportError("XLSX question has no exact writable location")
                    try:
                        sheet = workbook[sheet_name]
                    except KeyError as exc:
                        raise UnsafeExportError(
                            f"XLSX sheet location no longer exists: {sheet_name}"
                        ) from exc
                    if not isinstance(sheet, Worksheet):
                        raise UnsafeExportError("XLSX target is not a worksheet")
                    if sheet.sheet_state != "visible":
                        raise UnsafeExportError("XLSX question sheet is hidden")
                    target = _xlsx_target(sheet, cell_name)
                    target.value = neutralize_spreadsheet_formula(
                        _final_text(answer, reviews.get(answer.id))
                    )
                workbook.save(temporary)
            finally:
                workbook.close()
            _verify_source_identity(questionnaire, source)

    def _docx(
        self,
        questionnaire: Questionnaire,
        answers: dict[str, DraftAnswer],
        reviews: dict[str, ReviewDecision],
        source: Path,
        destination: Path,
    ) -> None:
        document = Document(str(source))
        for question in questionnaire.questions:
            answer = answers[question.id]
            text = _final_text(answer, reviews.get(answer.id))
            if question.location.paragraph is not None:
                try:
                    paragraph = document.paragraphs[question.location.paragraph]
                except IndexError as exc:
                    raise UnsafeExportError("DOCX paragraph location no longer exists") from exc
            elif question.location.key:
                paragraph = _docx_paragraph(document, question.location.key)
            else:
                raise UnsafeExportError("DOCX question has no writable location")
            paragraph.add_run(f"\nAnswer: {text}")
        with _atomic_destination(destination) as temporary:
            document.save(str(temporary))
            _verify_source_identity(questionnaire, source)

    def _csv(
        self,
        questionnaire: Questionnaire,
        answers: dict[str, DraftAnswer],
        reviews: dict[str, ReviewDecision],
        source: Path,
        destination: Path,
    ) -> None:
        with source.open(newline="", encoding="utf-8-sig") as handle:
            rows = list(csv.reader(handle))
        for question in questionnaire.questions:
            answer = answers[question.id]
            row = (question.location.row or 1) - 1
            try:
                question_column = int(question.location.key or "0")
            except ValueError as exc:
                raise UnsafeExportError("CSV question column is invalid") from exc
            if question_column < 0:
                raise UnsafeExportError("CSV question column is invalid")
            column = question_column + 1
            if row < 0 or row >= len(rows):
                raise UnsafeExportError("CSV question row no longer exists")
            if len(rows[row]) > column and rows[row][column] != "":
                raise UnsafeExportError(
                    f"CSV answer target is occupied at row {row + 1}, column {column + 1}"
                )
            while len(rows[row]) <= column:
                rows[row].append("")
            rows[row][column] = neutralize_spreadsheet_formula(
                _final_text(answer, reviews.get(answer.id))
            )
        with (
            _atomic_destination(destination) as temporary,
            temporary.open("w", newline="", encoding="utf-8") as handle,
        ):
            csv.writer(handle).writerows(rows)
            handle.flush()
            os.fsync(handle.fileno())
            _verify_source_identity(questionnaire, source)

    def _json(
        self,
        questionnaire: Questionnaire,
        answers: dict[str, DraftAnswer],
        reviews: dict[str, ReviewDecision],
        source: Path,
        destination: Path,
    ) -> None:
        payload = []
        for question in questionnaire.questions:
            answer = answers[question.id]
            review = reviews.get(answer.id)
            payload.append(
                {
                    "question_id": question.id,
                    "question": question.text,
                    "answer": _final_text(answer, review),
                    "draft_status": answer.status.value,
                    "review_state": review.state.value if review is not None else None,
                    "sources": [
                        {
                            "id": item.source_id,
                            "version": item.source_version,
                            "digest": item.source_digest,
                            "uri": item.source_uri,
                        }
                        for item in answer.evidence
                    ],
                }
            )
        with _atomic_destination(destination) as temporary:
            temporary.write_text(
                json.dumps(payload, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            _verify_source_identity(questionnaire, source)
