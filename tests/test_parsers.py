import csv
import json

from docx import Document
from openpyxl import Workbook

from trustflow.adapters.parsers import ParserRegistry
from trustflow.domain.models import DocumentFormat, QuestionSensitivity


def test_json_parser(tmp_path) -> None:
    path = tmp_path / "q.json"
    path.write_text(json.dumps({"questions": ["Do you encrypt data?"]}), encoding="utf-8")
    questionnaire = ParserRegistry().parse(path)
    assert questionnaire.format is DocumentFormat.JSON
    assert questionnaire.questions[0].sensitivity is QuestionSensitivity.SECURITY


def test_markdown_parser(tmp_path) -> None:
    path = tmp_path / "q.md"
    path.write_text("# Do you encrypt data?\nNot a question\n", encoding="utf-8")
    questionnaire = ParserRegistry().parse(path)
    assert len(questionnaire.questions) == 1


def test_csv_parser(tmp_path) -> None:
    path = tmp_path / "q.csv"
    with path.open("w", newline="", encoding="utf-8") as handle:
        csv.writer(handle).writerow(["Do you encrypt data?", ""])
    questionnaire = ParserRegistry().parse(path)
    assert questionnaire.questions[0].location.row == 1


def test_xlsx_parser(tmp_path) -> None:
    path = tmp_path / "q.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "Do you encrypt data?"
    workbook.save(path)
    questionnaire = ParserRegistry().parse(path)
    assert questionnaire.questions[0].location.cell == "A1"


def test_docx_parser(tmp_path) -> None:
    path = tmp_path / "q.docx"
    document = Document()
    document.add_paragraph("Do you encrypt data?")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Do you have legal insurance?"
    document.save(path)
    questionnaire = ParserRegistry().parse(path)
    assert len(questionnaire.questions) == 2
