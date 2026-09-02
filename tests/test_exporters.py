import csv
import json
from datetime import UTC, datetime

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from trustflow.adapters.exporters import ExporterRegistry
from trustflow.adapters.parsers import ParserRegistry
from trustflow.domain.errors import InvalidTransitionError, UnsafeExportError
from trustflow.domain.models import AnswerStatus, DraftAnswer, Evidence


def evidence(*, source_digest: str = "a" * 64) -> Evidence:
    return Evidence(
        source_id="source",
        source_title="Source",
        source_uri="policy://source",
        source_version="1",
        source_digest=source_digest,
        owner="owner",
        excerpt="Approved evidence.",
        score=1,
        updated_at=datetime.now(UTC),
    )


def answer(questionnaire_id, question_id, text) -> DraftAnswer:
    return DraftAnswer(
        questionnaire_id=questionnaire_id,
        question_id=question_id,
        text=text,
        status=AnswerStatus.ANSWERED,
        confidence=1,
        evidence=(evidence(),),
    )


def test_xlsx_export_neutralizes_formula(tmp_path) -> None:
    source = tmp_path / "q.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "Question?"
    workbook.save(source)
    workbook.close()
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.xlsx"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", '=HYPERLINK("bad")')],
        {},
        output,
    )
    reopened = load_workbook(output)
    try:
        assert reopened.active["B1"].value.startswith("'")
    finally:
        reopened.close()


def test_xlsx_export_rejects_chartsheet_target(tmp_path) -> None:
    from openpyxl.chart import BarChart, Reference

    source = tmp_path / "chartsheet.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    assert worksheet is not None
    worksheet["A1"] = "Question?"
    worksheet["B1"] = 1
    worksheet["B2"] = 2
    chartsheet = workbook.create_chartsheet("Chart")
    chart = BarChart()
    chart.add_data(Reference(worksheet, min_col=2, min_row=1, max_row=2))
    chartsheet.add_chart(chart)
    workbook.save(source)
    workbook.close()

    questionnaire = ParserRegistry().parse(source)
    question = questionnaire.questions[0]
    questionnaire = questionnaire.model_copy(
        update={
            "questions": (
                question.model_copy(
                    update={
                        "location": question.location.model_copy(
                            update={"sheet": "Chart", "cell": "A1"}
                        )
                    }
                ),
            )
        }
    )

    with pytest.raises(UnsafeExportError, match="not a worksheet"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, question.id, "Answer")],
            {},
            tmp_path / "out.xlsx",
        )


def test_xlsx_export_refuses_occupied_adjacent_cell(tmp_path) -> None:
    source = tmp_path / "occupied.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "Question?"
    sheet["B1"] = "Existing answer"
    workbook.save(source)
    workbook.close()
    questionnaire = ParserRegistry().parse(source)

    with pytest.raises(UnsafeExportError, match="answer target is occupied"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Replacement")],
            {},
            tmp_path / "out.xlsx",
        )


def test_xlsx_export_places_answer_after_merged_question_range(tmp_path) -> None:
    source = tmp_path / "merged-question.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.merge_cells("A1:B1")
    sheet["A1"] = "Merged question?"
    workbook.save(source)
    workbook.close()
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.xlsx"

    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", "Answer")],
        {},
        output,
    )

    reopened = load_workbook(output)
    try:
        assert reopened.active["C1"].value == "Answer"
    finally:
        reopened.close()


def test_xlsx_export_rejects_merged_answer_target(tmp_path) -> None:
    source = tmp_path / "merged-target.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "Question?"
    sheet.merge_cells("B1:C1")
    workbook.save(source)
    workbook.close()
    questionnaire = ParserRegistry().parse(source)

    with pytest.raises(UnsafeExportError, match="merged range"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Answer")],
            {},
            tmp_path / "out.xlsx",
        )


def test_csv_export(tmp_path) -> None:
    source = tmp_path / "q.csv"
    with source.open("w", newline="", encoding="utf-8") as handle:
        csv.writer(handle).writerow(["Question?"])
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.csv"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", "Answer")],
        {},
        output,
    )
    assert "Answer" in output.read_text(encoding="utf-8")


def test_csv_export_refuses_occupied_adjacent_cell(tmp_path) -> None:
    source = tmp_path / "occupied.csv"
    with source.open("w", newline="", encoding="utf-8") as handle:
        csv.writer(handle).writerow(["Question?", "Existing answer"])
    questionnaire = ParserRegistry().parse(source)

    with pytest.raises(UnsafeExportError, match="answer target is occupied"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Replacement")],
            {},
            tmp_path / "out.csv",
        )


def test_docx_export(tmp_path) -> None:
    source = tmp_path / "q.docx"
    document = Document()
    document.add_paragraph("Question?")
    document.save(source)
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.docx"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", "Answer")],
        {},
        output,
    )
    assert "Answer" in Document(output).paragraphs[0].text


def test_json_export(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.json"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", "Answer")],
        {},
        output,
    )
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload[0]["answer"] == "Answer"
    assert payload[0]["sources"][0]["digest"] == "a" * 64


def test_formula_neutralization_with_leading_whitespace(tmp_path) -> None:
    source = tmp_path / "q.csv"
    source.write_text("Question?\n", encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    output = tmp_path / "out.csv"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", '\t=HYPERLINK("bad")')],
        {},
        output,
    )
    assert "'\t=HYPERLINK" in output.read_text(encoding="utf-8")


def test_export_refuses_to_overwrite_source(tmp_path) -> None:
    source = tmp_path / "q.json"
    original = '{"questions":["Question?"]}'
    source.write_text(original, encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    with pytest.raises(UnsafeExportError, match="must differ"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Answer")],
            {},
            source,
        )
    assert source.read_text(encoding="utf-8") == original


def test_export_refuses_to_overwrite_existing_destination(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    destination = tmp_path / "existing.json"
    destination.write_text("do not replace", encoding="utf-8")

    with pytest.raises(UnsafeExportError, match="already exists"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Answer")],
            {},
            destination,
        )
    assert destination.read_text(encoding="utf-8") == "do not replace"


def test_export_blocks_questionnaire_changed_since_import(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    source.write_text('{"questions":["Changed question?"]}', encoding="utf-8")

    with pytest.raises(UnsafeExportError, match="changed since import"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Answer")],
            {},
            tmp_path / "out.json",
        )


def test_export_blocks_questionnaire_without_source_fingerprint(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source).model_copy(update={"source_digest": "0" * 64})

    with pytest.raises(UnsafeExportError, match="source fingerprint is missing"):
        ExporterRegistry().export(
            questionnaire,
            [answer(questionnaire.id, "q1", "Answer")],
            {},
            tmp_path / "out.json",
        )


def test_docx_nested_table_round_trip(tmp_path) -> None:
    source = tmp_path / "nested.docx"
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).paragraphs[0].text = "Nested question?"
    document.save(source)
    questionnaire = ParserRegistry().parse(source)
    assert len(questionnaire.questions) == 1
    output = tmp_path / "nested-out.docx"
    ExporterRegistry().export(
        questionnaire,
        [answer(questionnaire.id, "q1", "Nested answer")],
        {},
        output,
    )
    reopened = Document(output)
    nested_output = reopened.tables[0].cell(0, 0).tables[0].cell(0, 0).text
    assert "Nested answer" in nested_output


def test_exporter_defense_in_depth_blocks_unreviewed_status(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Sensitive question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    unresolved = DraftAnswer(
        questionnaire_id=questionnaire.id,
        question_id="q1",
        text="Draft",
        status=AnswerStatus.REVIEW_REQUIRED,
        confidence=0.8,
        evidence=(evidence(),),
    )
    with pytest.raises(InvalidTransitionError, match="unresolved"):
        ExporterRegistry().export(
            questionnaire,
            [unresolved],
            {},
            tmp_path / "out.json",
        )


def test_exporter_defense_in_depth_blocks_claim_without_evidence(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    unsupported = DraftAnswer(
        questionnaire_id=questionnaire.id,
        question_id="q1",
        text="Unsupported claim",
        status=AnswerStatus.ANSWERED,
        confidence=1,
    )
    with pytest.raises(InvalidTransitionError, match="without evidence"):
        ExporterRegistry().export(
            questionnaire,
            [unsupported],
            {},
            tmp_path / "out.json",
        )


def test_exporter_blocks_evidence_without_source_fingerprint(tmp_path) -> None:
    source = tmp_path / "q.json"
    source.write_text('{"questions":["Question?"]}', encoding="utf-8")
    questionnaire = ParserRegistry().parse(source)
    legacy = DraftAnswer(
        questionnaire_id=questionnaire.id,
        question_id="q1",
        text="Legacy claim",
        status=AnswerStatus.ANSWERED,
        confidence=1,
        evidence=(evidence(source_digest="0" * 64),),
    )
    with pytest.raises(InvalidTransitionError, match="missing a source fingerprint"):
        ExporterRegistry().export(
            questionnaire,
            [legacy],
            {},
            tmp_path / "out.json",
        )
